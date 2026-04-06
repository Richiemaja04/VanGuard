"""
utils.py — Helper Functions for ResumeIQ
"""

import uuid
import aiofiles
from pathlib import Path
from fastapi import UploadFile, HTTPException

UPLOAD_DIR = Path("static/uploads")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc"}
MAX_FILE_SIZE_MB = 10


# -------------------------------
# FILE VALIDATION
# -------------------------------
def allowed_file(filename: str) -> bool:
    if not filename:
        return False
    return Path(filename).suffix.lower() in ALLOWED_EXTENSIONS


# -------------------------------
# SAVE UPLOAD
# -------------------------------
async def save_upload(file: UploadFile) -> str:
    if not file.filename:
        raise HTTPException(status_code=400, detail="File must have a valid name")

    ext = Path(file.filename).suffix.lower()

    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="Invalid file type")

    unique_name = f"{uuid.uuid4().hex}{ext}"
    save_path = UPLOAD_DIR / unique_name

    content = await file.read()

    # File size check
    size_mb = len(content) / (1024 * 1024)
    if size_mb > MAX_FILE_SIZE_MB:
        raise HTTPException(
            status_code=400,
            detail=f"File too large: {size_mb:.1f}MB (max {MAX_FILE_SIZE_MB}MB)"
        )

    async with aiofiles.open(save_path, "wb") as f:
        await f.write(content)

    return str(save_path)


# -------------------------------
# HTML → PDF
# -------------------------------
def html_to_pdf(html_path: str, output_path: str) -> str:
    try:
        from weasyprint import HTML
        HTML(filename=html_path).write_pdf(output_path)
    except ImportError:
        import pdfkit
        pdfkit.from_file(html_path, output_path)

    return output_path


# -------------------------------
# HTML → DOCX
# -------------------------------
def html_to_docx(html_path: str, output_path: str) -> str:
    try:
        from htmldocx import HtmlToDocx
        from docx import Document

        with open(html_path, "r", encoding="utf-8") as f:
            html_content = f.read()

        doc = Document()
        parser = HtmlToDocx()
        parser.add_html_to_document(html_content, doc)
        doc.save(output_path)

    except ImportError:
        from docx import Document
        import re

        with open(html_path, "r", encoding="utf-8") as f:
            html_content = f.read()

        text = re.sub(r"<[^>]+>", " ", html_content)
        text = re.sub(r"\s+", " ", text).strip()

        doc = Document()
        for line in text.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

        doc.save(output_path)

    return output_path


# -------------------------------
# TEXT CLEANING
# -------------------------------
def clean_text(text: str) -> str:
    import re
    text = re.sub(r"\r\n", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# -------------------------------
# TEXT TRUNCATION
# -------------------------------
def truncate_text(text: str, max_chars: int = 4000) -> str:
    return text[:max_chars] if len(text) > max_chars else text